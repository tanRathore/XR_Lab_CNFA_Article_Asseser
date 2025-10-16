# -*- coding: utf-8 -*-
"""
Specialists: bibliography-first + thick synthesis with strict expansion.
- Robust JSON extraction (no recursive regex)
- Topic-agnostic reference categories (relevance over year)
- ≥30 references (augments once if short), DOI/URL encouraged
- 2,800–3,400 words target; hard floors per section
- 10–14 row empirical summary table
- Auto-expands thin drafts until thresholds met (one retry pass)
- Safe DOCX writer that renders markdown table
- Exports ALL agents expected by the router
"""

import sys
from pathlib import Path
import json
import re
from typing import Any, Dict, List, Optional, Tuple

# --- Setup Project Path ---
SRC_DIR = Path(__file__).resolve().parent.parent.parent
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))
# --- End Setup Project Path ---

# Third-party
import docx
from docx.shared import Inches  # noqa: F401

from langchain.agents import AgentExecutor, Tool, create_tool_calling_agent
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnableLambda
from langchain_experimental.agents.agent_toolkits import create_python_agent as _create_python_agent  # optional
from langchain_community.tools.tavily_search import TavilySearchResults

# Project
from cna_rag_agent.utils.logging_config import logger
from cna_rag_agent.generation.generator import get_llm
from cna_rag_agent.config import GEMINI_PRO_MODEL_NAME
from cna_rag_agent.vector_store.store_manager import get_retriever

# Optional tools
try:
    from cna_rag_agent.agent.advanced_tools import python_repl_tool, image_generator_tool
    _HAS_ADV_TOOLS = True
except Exception as e:
    logger.warning(f"[specialists] advanced_tools unavailable: {e}")
    _HAS_ADV_TOOLS = False


# --------------------------- GLOBAL TARGETS ---------------------------

MIN_REFS = 30  # bibliography size target (>=)
TARGET_WORDS_MIN = 2800
TARGET_WORDS_MAX = 3400

# Section floors to prevent “thin” outputs
FLOOR_INTRO = 350
FLOOR_THEORY_SECTIONS = 3
FLOOR_THEORY_EACH = 250
FLOOR_KEYFINDINGS = 700
FLOOR_CONCLUSION = 300
FLOOR_TABLE_ROWS_MIN = 10
FLOOR_TABLE_ROWS_MAX = 14


# --------------------------- UTILITIES ---------------------------

def _pretty(obj: Any) -> str:
    try:
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except Exception:
        return str(obj)

def _extract_json_any(text: str) -> Optional[Dict[str, Any]]:
    """
    Robustly extract a top-level JSON object from a model response.
    Tries:
      1) Direct JSON
      2) ```json fenced block
      3) Outermost brace match with depth counting
    """
    if not text:
        return None

    s = text.strip()

    # 1) Direct
    if s.startswith("{") and s.endswith("}"):
        try:
            return json.loads(s)
        except Exception:
            pass

    # 2) Fenced
    m = re.search(r"```json\s*(\{.*?\})\s*```", text, flags=re.DOTALL | re.IGNORECASE)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception:
            pass

    # 3) Outermost braces (depth-based)
    first = text.find("{")
    last = text.rfind("}")
    if first != -1 and last != -1 and last > first:
        body = text[first:last+1]
        depth = 0
        end = -1
        for i, ch in enumerate(body):
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    end = i
                    break
        if end != -1:
            candidate = body[:end+1]
            try:
                return json.loads(candidate)
            except Exception:
                candidate2 = re.sub(r",\s*}", "}", candidate)
                try:
                    return json.loads(candidate2)
                except Exception:
                    return None
    return None


def _dedupe_refs(refs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = set()
    out = []
    for r in refs:
        key = (r.get("full_citation") or "").strip().lower()
        if key and key not in seen:
            seen.add(key)
            out.append(r)
    return out


def _validate_biblio(refs: List[Dict[str, Any]]) -> Tuple[bool, str]:
    if not isinstance(refs, list):
        return False, "Bibliography is not a list."
    if len(refs) < MIN_REFS:
        return False, f"Only {len(refs)} references; need ≥ {MIN_REFS}."
    for i, r in enumerate(refs):
        if not isinstance(r, dict) or not r.get("full_citation"):
            return False, f"Ref {i} missing full_citation."
    return True, "ok"


def _count_words(text: str) -> int:
    return len(re.findall(r"\w+", text or ""))


def _count_table_rows(md: str) -> int:
    if not md or "|" not in md:
        return 0
    lines = [ln for ln in (md.strip().splitlines()) if "|" in ln.strip()]
    if len(lines) < 2:
        return 0
    # naive: rows after header (& separator)
    if len(lines) >= 2 and set(lines[1].replace("|", "").strip()) <= set("-:"):
        return max(0, len(lines) - 2)
    return max(0, len(lines) - 1)


# --------------------------- DOCX WRITER ---------------------------

def _parse_markdown_table_to_docx(doc: docx.Document, table_md: str) -> None:
    if not table_md or '|' not in table_md:
        return
    lines = [ln.strip() for ln in table_md.strip().splitlines() if '|' in ln]
    if len(lines) < 2:
        return
    header = [h.strip() for h in lines[0].strip('|').split('|')]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
    # skip separator row if present
    start = 2 if len(lines) >= 2 and set(lines[1].replace("|", "").strip()) <= set("-:") else 1
    for row_line in lines[start:]:
        cols = [c.strip() for c in row_line.strip('|').split('|')]
        row = table.add_row().cells
        for i, c in enumerate(cols[:len(row)]):
            row[i].text = c


def save_report_to_docx(report_data: dict, filename: str = "Research_Report.docx"):
    try:
        doc = docx.Document()
        doc.add_heading(report_data.get('title', 'Research Report'), level=1)

        # Introduction
        doc.add_heading('Introduction', level=2)
        doc.add_paragraph(report_data.get('introduction', '').strip())

        # Theoretical Foundations
        doc.add_heading('Theoretical Foundations', level=2)
        tf = report_data.get('theoretical_foundations') or report_data.get('theoretical_foundations_text')
        if isinstance(tf, list):
            for item in tf:
                head = item.get("heading") or ""
                cnt = item.get("content") or ""
                if head:
                    doc.add_heading(head, level=3)
                if cnt:
                    for para in str(cnt).split("\n\n"):
                        if para.strip():
                            doc.add_paragraph(para.strip())
        elif isinstance(tf, str):
            for para in tf.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

        # Key Findings
        doc.add_heading('Key Findings and Methodologies', level=2)
        for para in (report_data.get('key_findings_and_methodologies', '') or '').split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Table
        doc.add_heading('Summary Table of Empirical Findings', level=2)
        table_md = report_data.get('summary_table_markdown') or report_data.get('summary_table') or ''
        _parse_markdown_table_to_docx(doc, table_md)

        # Conclusion
        doc.add_heading('Conclusion and Future Directions', level=2)
        for para in (report_data.get('conclusion_and_future_directions', '') or '').split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # References
        doc.add_heading('References', level=2)
        refs = report_data.get('references') or []
        if isinstance(refs, list):
            for ref in refs:
                p = doc.add_paragraph(style='List Bullet')
                cat = (ref.get('category') or '').strip()
                if cat:
                    r = p.add_run(f"({cat}) ")
                    r.bold = True
                p.add_run(ref.get('full_citation', '').strip())
                extras = []
                if ref.get("doi"):
                    extras.append(f"doi:{ref['doi']}")
                if ref.get("url"):
                    extras.append(ref["url"])
                if extras:
                    doc.add_paragraph(" | ".join(extras))

        doc.save(filename)
        logger.info(f"[save_report_to_docx] Saved report to '{filename}'")
    except Exception as e:
        logger.error(f"[save_report_to_docx] Failed to save DOCX: {e}", exc_info=True)


# --------------------------- LEAD RESEARCHER AGENT ---------------------------

def create_professional_researcher_agent() -> AgentExecutor:
    logger.info("Creating the 'Gold Standard' Professional Researcher AGENT...")

    # Local KB search
    local_retriever = get_retriever(k_results=12)
    local_search_tool = Tool(
        name="local_knowledge_base_search",
        func=lambda q: [doc.page_content for doc in local_retriever.invoke(q)],
        description="Search local CNfA papers/handbook: definitions, paradigms, methods."
    )

    # Web search (Tavily)
    web_search_tool = TavilySearchResults(max_results=20, name="online_web_search")
    web_search_tool.description = "Find relevant scholarly articles; return titles and URLs."

    tools = [local_search_tool, web_search_tool]

    system_prompt = """You are a Distinguished Professor in Cognitive Neuroscience for Architecture (CNfA).
Your job is to harvest references broadly (local + web) and then synthesize a long-form, structured, academic report.

When asked for structured output, reply with ONLY valid JSON (no backticks, no commentary).
Relevance of references is key; publication year is not a constraint unless specified."""

    llm = get_llm(model_name=GEMINI_PRO_MODEL_NAME, temperature=0.1, max_output_tokens=8192)

    agent = create_tool_calling_agent(
        llm,
        tools,
        ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            # We allow chat history; internal calls provide [] if not supplied.
            MessagesPlaceholder("chat_history"),
            ("human", "{input}"),
            ("placeholder", "{agent_scratchpad}")
        ])
    )
    return AgentExecutor(agent=agent, tools=tools, verbose=True, handle_parsing_errors=True)


# Helper to invoke AgentExecutor with optional chat history
def _invoke_agent(agent: AgentExecutor, text: str, chat_history: Optional[List[Any]] = None) -> str:
    payload = {"input": text, "chat_history": chat_history or []}
    res = agent.invoke(payload)
    if isinstance(res, str):
        return res
    return res.get("output") or res.get("result") or _pretty(res)


# --------------------------- PROMPTS ---------------------------

def prompt_bibliography(topic: str) -> str:
    return f"""
Build a topic-appropriate, categorized bibliography for:

{topic}

Use tools as needed until you have a rich set.

Return ONLY valid JSON with this schema:

{{
  "references": [
    {{
      "category": "freeform label (e.g., Core Theory, Methods, Constructs, Environmental Complexity, Landmarks, Individual Differences, VR vs Real, Neuro/Imaging, Applied/Wayfinding, Developmental, Reviews/Meta, Debates, etc.)",
      "full_citation": "APA-style citation (authors, year, title, venue, volume(issue), pages)",
      "year": 2012,
      "doi": "10.xxxx/xxxxx" or null,
      "url": "https://..." or null
    }}
  ]
}}

Requirements:
- Include AT LEAST {MIN_REFS} relevant references (more is fine).
- Prioritize relevance over year (classics + contemporary are welcome).
- Populate multiple meaningful categories.
- Prefer DOIs and URLs when available.
- Respond with ONLY the JSON object (no prose).
"""


def prompt_bibliography_augment(topic: str, existing_refs_json: str) -> str:
    return f"""
Augment the existing bibliography for:

{topic}

EXISTING_REFS_JSON:
{existing_refs_json}

Return ONLY valid JSON with the SAME SCHEMA (single object with "references": [...]).
- Keep all existing items; add more to reach ≥ {MIN_REFS}; remove duplicates by full_citation.
- Maintain multiple categories; keep relevance over year; add DOIs/URLs when available.
"""


def prompt_synthesis(topic: str, references_json: str) -> str:
    return f"""
Write a comprehensive literature review grounded in the bibliography below.

TOPIC:
{topic}

BIBLIOGRAPHY_JSON:
{references_json}

Return ONLY valid JSON with EXACT schema:

{{
  "title": "string",
  "introduction": "multi-paragraph text",
  "theoretical_foundations": [
    {{"heading": "string", "content": "multi-paragraph text"}},
    {{"heading": "string", "content": "multi-paragraph text"}},
    {{"heading": "string", "content": "multi-paragraph text"}}
  ],
  "key_findings_and_methodologies": "multi-paragraph text covering paradigms, measures, and specific result patterns",
  "summary_table_markdown": "| Research Domain | Key Finding | Supporting Studies | Theoretical Implications |\\n|---|---|---|---|\\n| ... | ... | ... | ... |",
  "conclusion_and_future_directions": "multi-paragraph text with concrete future directions",
  "references": [
    {{
      "category": "string",
      "full_citation": "string",
      "year": 2019,
      "doi": "string or null",
      "url": "string or null"
    }}
  ]
}}

Quality constraints (critical):
- Overall length: ~{TARGET_WORDS_MIN}-{TARGET_WORDS_MAX} words.
- Introduction ≥ {FLOOR_INTRO} words.
- Theoretical Foundations: ≥ {FLOOR_THEORY_SECTIONS} subsections, each ≥ {FLOOR_THEORY_EACH} words.
- Key Findings & Methodologies ≥ {FLOOR_KEYFINDINGS} words, dense and specific.
- Conclusion ≥ {FLOOR_CONCLUSION} words with explicit, testable future work.
- Summary table: {FLOOR_TABLE_ROWS_MIN}–{FLOOR_TABLE_ROWS_MAX} rows with substantive cells.
- Write full content (no placeholders). Do NOT include inline citations—only the reference list at the end.
- Respond with ONLY the JSON object.
"""


def prompt_expand(current_json: str) -> str:
    return f"""
The following JSON report is too thin in one or more sections. EXPAND it to meet ALL constraints.

CURRENT_REPORT_JSON:
{current_json}

Return ONLY valid JSON with the SAME schema. Rules:
- Raise total length toward {TARGET_WORDS_MIN}-{TARGET_WORDS_MAX} words.
- Ensure ≥ {FLOOR_INTRO} words in Introduction.
- Ensure ≥ {FLOOR_THEORY_SECTIONS} theory subsections, each ≥ {FLOOR_THEORY_EACH} words.
- Ensure Key Findings & Methodologies ≥ {FLOOR_KEYFINDINGS} words.
- Ensure Conclusion ≥ {FLOOR_CONCLUSION} words.
- Ensure the summary table has {FLOOR_TABLE_ROWS_MIN}–{FLOOR_TABLE_ROWS_MAX} rows with substantive cells.
- Keep the full reference list (unchanged except minor formatting fixes). Respond ONLY with JSON.
"""


# --------------------------- PIPELINE (harvest → synthesize → expand) ---------------------------

def harvest_bibliography(agent: AgentExecutor, topic: str, chat_history: Optional[List[Any]] = None) -> List[Dict[str, Any]]:
    logger.info("[harvest_bibliography] start")
    raw = _invoke_agent(agent, prompt_bibliography(topic), chat_history)
    data = _extract_json_any(raw) or {"references": []}
    refs = _dedupe_refs(data.get("references", []))

    if len(refs) < MIN_REFS:
        logger.info(f"[harvest_bibliography] {len(refs)} refs; augmenting…")
        aug_raw = _invoke_agent(
            agent,
            prompt_bibliography_augment(topic, json.dumps({"references": refs}, ensure_ascii=False)),
            chat_history
        )
        data2 = _extract_json_any(aug_raw)
        if data2 and isinstance(data2.get("references"), list):
            refs = _dedupe_refs(refs + data2["references"])

    ok, msg = _validate_biblio(refs)
    if not ok:
        logger.warning(f"[harvest_bibliography] insufficient bibliography: {msg}")
    else:
        logger.info(f"[harvest_bibliography] bibliography ready: n={len(refs)}")
    return refs


def compose_report(agent: AgentExecutor, topic: str, refs: List[Dict[str, Any]], chat_history: Optional[List[Any]] = None) -> Dict[str, Any]:
    logger.info("[compose_report] synthesize")
    refs_json = json.dumps({"references": refs}, ensure_ascii=False)
    raw = _invoke_agent(agent, prompt_synthesis(topic, refs_json), chat_history)
    data = _extract_json_any(raw)

    if not data:
        logger.warning("[compose_report] parse failed; second attempt")
        raw2 = _invoke_agent(agent, prompt_synthesis(topic, refs_json), chat_history)
        data = _extract_json_any(raw2)

    if not data:
        logger.error("[compose_report] still failed; fallback skeleton")
        return {
            "title": "Research Report",
            "introduction": "",
            "theoretical_foundations": [],
            "key_findings_and_methodologies": "",
            "summary_table_markdown": "| Research Domain | Key Finding | Supporting Studies | Theoretical Implications |\n|---|---|---|---|",
            "conclusion_and_future_directions": "",
            "references": refs
        }

    # Ensure references present
    if not data.get("references"):
        data["references"] = refs

    return data


def _needs_expansion(report: Dict[str, Any]) -> bool:
    intro_wc = _count_words(report.get("introduction", ""))
    theory = report.get("theoretical_foundations", []) or []
    key_wc = _count_words(report.get("key_findings_and_methodologies", ""))
    concl_wc = _count_words(report.get("conclusion_and_future_directions", ""))
    table_rows = _count_table_rows(report.get("summary_table_markdown", "") or report.get("summary_table", ""))

    conds = [
        intro_wc < FLOOR_INTRO,
        len(theory) < FLOOR_THEORY_SECTIONS,
        any(_count_words(sec.get("content", "")) < FLOOR_THEORY_EACH for sec in theory[:FLOOR_THEORY_SECTIONS]) if theory else True,
        key_wc < FLOOR_KEYFINDINGS,
        concl_wc < FLOOR_CONCLUSION,
        table_rows < FLOOR_TABLE_ROWS_MIN or table_rows > FLOOR_TABLE_ROWS_MAX
    ]
    return any(conds)


def expand_if_thin(agent: AgentExecutor, report: Dict[str, Any], chat_history: Optional[List[Any]] = None) -> Dict[str, Any]:
    if not _needs_expansion(report):
        return report
    logger.info("[expand_if_thin] expanding thin report")
    raw = _invoke_agent(agent, prompt_expand(json.dumps(report, ensure_ascii=False)), chat_history)
    expanded = _extract_json_any(raw)
    if expanded and not _needs_expansion(expanded):
        return expanded
    logger.warning("[expand_if_thin] expansion insufficient; returning best-so-far")
    return expanded or report


# Convenience: full pipeline for routing
def _run_report_pipeline(payload: Dict[str, Any]) -> Dict[str, str]:
    """
    Runnable-compatible wrapper:
      payload: { 'input': topic, 'chat_history': [...] }
      returns: { 'output': JSON-string-of-report }
    """
    topic = (payload.get("input") or "").strip() or "Unspecified topic"
    chat_history = payload.get("chat_history") or []
    lead = create_professional_researcher_agent()
    refs = harvest_bibliography(lead, topic, chat_history)
    report = compose_report(lead, topic, refs, chat_history)
    report = expand_if_thin(lead, report, chat_history)
    return {"output": json.dumps(report, ensure_ascii=False)}


# Expose a pipeline Runnable for the router (optional use in slash path)
lead_researcher_pipeline = RunnableLambda(_run_report_pipeline)


# --------------------------- SIMPLE SPECIALISTS (for router) ---------------------------

def _make_simple_agent(system_prompt: str, temperature: float = 0.2):
    llm = get_llm(model_name=GEMINI_PRO_MODEL_NAME, temperature=temperature)
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}")
    ])
    chain = prompt | llm | StrOutputParser()
    return (chain | RunnableLambda(lambda s: {"output": s}))


def _make_data_scientist_agent():
    if _HAS_ADV_TOOLS:
        try:
            llm = get_llm(model_name=GEMINI_PRO_MODEL_NAME, temperature=0.0)
            py_agent = _create_python_agent(llm=llm, tool=python_repl_tool, verbose=True)
            def _call(payload: Dict[str, Any]):
                res = py_agent.invoke({"input": payload.get("input", "")})
                if isinstance(res, dict) and "output" in res:
                    return {"output": res["output"]}
                return {"output": str(res)}
            return RunnableLambda(_call)
        except Exception as e:
            logger.warning(f"[specialists] python agent fallback due to: {e}")
    return _make_simple_agent(
        "You are a rigorous data scientist. Show formulas, step-by-step calculations, and double-check results. "
        "If code is needed, write and explain it.", temperature=0.0
    )


def _make_visualizer_agent():
    if _HAS_ADV_TOOLS:
        def _invoke(payload: Dict[str, Any]) -> Dict[str, str]:
            text = payload.get("input") or ""
            try:
                result = image_generator_tool.invoke({"prompt": text})
                return {"output": str(result)}
            except Exception as e:
                return {"output": f"(Visualizer tool failed) {e}"}
        return RunnableLambda(_invoke)
    return _make_simple_agent(
        "You are a visual explainer. When asked to visualize, describe a clear, step-by-step diagram "
        "or figure that would communicate the concept well.", temperature=0.0
    )


# Instantiate the heavy agent (if you need direct access)
logger.info("Initializing Lead Researcher (tool-using) agent...")
lead_researcher_agent = create_professional_researcher_agent()
logger.info("Lead Researcher agent ready.")

# Other specialists used by the router
information_agent = _make_simple_agent(
    "Provide concise, high-quality overviews and primers with examples and definitions."
)
list_articles_agent = _make_simple_agent(
    "List relevant articles from the local knowledge and general domain. Group by theme, cite clearly."
)
deep_dive_agent = _make_simple_agent(
    "Explain detailed methods and experimental setups. Include measures, variables, procedures, and pitfalls."
)
open_questions_agent = _make_simple_agent(
    "Identify gaps and open questions. Propose concrete experiments and measurable predictions."
)
comparative_analyst_agent = _make_simple_agent(
    "Compare and contrast topics across theory, evidence, methods, pros/cons, and applicability."
)
online_researcher_agent = _make_simple_agent(
    "Summarize what the latest literature is likely to say and propose search strings and databases to query. "
    "Flag where fresh web verification would be needed."
)
general_qa_agent = _make_simple_agent(
    "Answer succinctly with a single authoritative definition or fact, and one pointer to read more."
)
data_scientist_agent = _make_data_scientist_agent()
visualizer_agent = _make_visualizer_agent()

__all__ = [
    # pipeline helpers
    "save_report_to_docx",
    "harvest_bibliography",
    "compose_report",
    "expand_if_thin",
    "lead_researcher_pipeline",
    # agents
    "lead_researcher_agent",
    "information_agent",
    "list_articles_agent",
    "deep_dive_agent",
    "open_questions_agent",
    "comparative_analyst_agent",
    "online_researcher_agent",
    "general_qa_agent",
    "data_scientist_agent",
    "visualizer_agent",
]
