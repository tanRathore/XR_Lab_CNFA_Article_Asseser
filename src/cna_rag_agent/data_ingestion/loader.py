# src/cna_rag_agent/data_ingestion/loader.py

import os
import re
import docx
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Iterator

from langchain_core.documents import Document
from langchain_community.document_loaders import UnstructuredFileLoader

try:
    from ..config import RAW_DOCUMENTS_PATH
    from ..utils.logging_config import logger
except ImportError:
    import sys
    SRC_DIR = Path(__file__).resolve().parent.parent.parent
    if str(SRC_DIR) not in sys.path: sys.path.insert(0, str(SRC_DIR))
    from cna_rag_agent.config import RAW_DOCUMENTS_PATH
    from cna_rag_agent.utils.logging_config import logger


# --- FINAL, CORRECTED PARSER FOR THE HANDBOOK ---
def parse_handbook_docx(file_path: Path) -> Iterator[Document]:
    """
    Parses the structured 'Handbook of CNfA Experimental Paradigms.docx'.
    It reads the document paragraph by paragraph, identifies hierarchical headings,
    and correctly handles content on the same line as a heading.
    """
    logger.info(f"Using final, corrected handbook parser for: {file_path.name}")
    
    category_regex = re.compile(r"^\s*([IVXLCDM]+)\.\s+(.*)")
    paradigm_regex = re.compile(r"^\s*(\??[A-Z])\.\s+(.*)")
    section_headings = [
        "What is Studied", "Why it Matters", "Procedural Summary Table",
        "Theoretical Background & Major Explanations", "Example Findings",
        "Explanation of Measurables", "Stimuli Description"
    ]

    doc = docx.Document(file_path)
    
    current_category = "Uncategorized"
    current_paradigm = "Uncategorized"
    current_section = "General"
    
    buffer = []

    def yield_buffer() -> Iterator[Document]:
        nonlocal buffer
        if buffer:
            content = "\n".join(buffer).strip()
            if content:
                metadata = {
                    "source": str(file_path.resolve()),
                    "file_name": file_path.name,
                    "category": current_category,
                    "paradigm_name": current_paradigm,
                    "paradigm_section": current_section,
                    "content_type": "handbook_entry"
                }
                yield Document(page_content=content, metadata=metadata)
            buffer = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        category_match = category_regex.match(text)
        paradigm_match = paradigm_regex.match(text)
        
        # --- NEW LOGIC TO HANDLE HEADINGS AND CONTENT ON THE SAME LINE ---
        found_heading = False
        
        if category_match:
            yield from yield_buffer()
            current_category = text
            current_paradigm = "General Introduction"
            current_section = "Category Title"
            found_heading = True
            # No content to extract from the category line itself
            
        elif paradigm_match:
            yield from yield_buffer()
            current_paradigm = text
            current_section = "Paradigm Title"
            found_heading = True
            # No content to extract from the paradigm line itself
            
        else:
            for heading in section_headings:
                if text.startswith(heading):
                    yield from yield_buffer()
                    current_section = heading
                    found_heading = True
                    # If there's content after the heading, add it to the new buffer
                    content_after_heading = text[len(heading):].strip()
                    # Remove leading dashes or colons often found after headings
                    if content_after_heading.startswith(('–', '—', ':')):
                        content_after_heading = content_after_heading[1:].strip()
                    
                    if content_after_heading:
                        buffer.append(content_after_heading)
                    break # Stop checking for other headings
        
        if not found_heading:
            # If the line is not a heading, it's regular content
            buffer.append(text)
        # --- END OF NEW LOGIC ---

    yield from yield_buffer() # Yield any remaining content


SUPPORTED_EXTENSIONS: Dict[str, Tuple[type[UnstructuredFileLoader], Dict]] = {
    ".pdf": (UnstructuredFileLoader, {"mode": "elements", "strategy": "hi_res"}),
    ".docx": (UnstructuredFileLoader, {"mode": "elements", "strategy": "fast"}),
    ".doc": (UnstructuredFileLoader, {"mode": "elements", "strategy": "fast"}),
    ".txt": (UnstructuredFileLoader, {"mode": "elements"}),
}

# The rest of the file remains unchanged.
def load_single_document(file_path: Path) -> List[Document]:
    file_ext = file_path.suffix.lower()
    
    logger.debug(f"Attempting to load document: {file_path} (extension: {file_ext})")
    if file_ext not in SUPPORTED_EXTENSIONS:
        logger.warning(f"Unsupported file extension '{file_ext}' for file: {file_path}. Skipping.")
        return []

    LoaderClass, loader_kwargs = SUPPORTED_EXTENSIONS[file_ext]
    loader = LoaderClass(str(file_path), **loader_kwargs)
    try:
        loaded_docs = loader.load()
        logger.info(f"Successfully loaded {len(loaded_docs)} element(s) from: {file_path}")
        for doc_element in loaded_docs:
            doc_element.metadata["file_name"] = file_path.name
            doc_element.metadata["file_path"] = str(file_path.resolve())
            if "source" not in doc_element.metadata: doc_element.metadata["source"] = str(file_path.resolve())
            if 'page_number' not in doc_element.metadata and 'page' in doc_element.metadata: doc_element.metadata['page_number'] = doc_element.metadata['page']
        return loaded_docs
    except Exception as e:
        logger.error(f"Error loading document {file_path}: {e}", exc_info=True)
        return []

def load_all_documents(documents_dir: Path = RAW_DOCUMENTS_PATH) -> List[Document]:
    if not documents_dir.exists() or not documents_dir.is_dir():
        logger.error(f"Documents directory '{documents_dir}' does not exist or is not a directory.")
        return []
    logger.info(f"Starting to load documents from directory: {documents_dir}")
    all_loaded_documents: List[Document] = []
    files_processed = 0
    files_skipped = 0
    
    for item in documents_dir.iterdir():
        if item.is_file():
            files_processed += 1
            if item.name == "BEST_Handbook of CNfA Experimental Paradigms.docx":
                docs_from_file = list(parse_handbook_docx(item))
                if docs_from_file:
                    all_loaded_documents.extend(docs_from_file)
                else:
                    logger.warning(f"Specialized parser for {item.name} yielded no documents.")
                    files_skipped += 1
            else:
                docs_from_file = load_single_document(item)
                if docs_from_file:
                    all_loaded_documents.extend(docs_from_file)
                else:
                    files_skipped += 1
        elif item.is_dir():
            logger.info(f"Found subdirectory '{item.name}'. Recursive loading not implemented. Skipping.")

    logger.info(f"Finished loading documents from {documents_dir}.")
    logger.info(f"Total files found (top-level): {files_processed}")
    logger.info(f"Files successfully processed into elements: {files_processed - files_skipped}")
    logger.info(f"Files skipped (unsupported or error): {files_skipped}")
    logger.info(f"Total document elements loaded: {len(all_loaded_documents)}")
    return all_loaded_documents


def get_full_text_for_article(file_path: Path) -> Optional[str]:
    logger.info(f"Attempting to load full text for article: {file_path}")
    document_elements = load_single_document(file_path)
    if not document_elements:
        logger.error(f"Failed to load any elements from article: {file_path}")
        return None
    full_text = "\n\n".join([doc.page_content for doc in document_elements if doc.page_content and doc.page_content.strip()])
    if not full_text.strip():
        logger.warning(f"Full text for article {file_path} is empty after joining elements.")
        return None
    logger.info(f"Successfully extracted full text for article: {file_path} (length: {len(full_text)} chars)")
    return full_text

if __name__ == "__main__":
    logger.info("Running loader.py directly for testing...")
    test_docs_path = RAW_DOCUMENTS_PATH
    
    if not any(f for f in test_docs_path.iterdir() if f.is_file() and not f.name.startswith('.') and f.suffix):
        dummy_txt_path = test_docs_path / "dummy_loader_test.txt"
        with open(dummy_txt_path, "w") as f:
            f.write("This is a simple test document for the loader direct run.")
        logger.info(f"Created dummy test file for loader.py: {dummy_txt_path}")
        
    loaded_documents = load_all_documents(test_docs_path)
    if loaded_documents:
        logger.info(f"load_all_documents test loaded {len(loaded_documents)} elements.")
        
        handbook_chunk_found = False
        for doc in loaded_documents:
            if doc.metadata.get("content_type") == "handbook_entry":
                logger.info("\n--- Example of a parsed handbook chunk ---")
                logger.info(f"Content: {doc.page_content[:200]}...")
                logger.info(f"Metadata: {doc.metadata}")
                handbook_chunk_found = True
                break
        if not handbook_chunk_found:
            logger.warning("Could not find a chunk parsed from the handbook to display. Ensure the file is in the raw_documents folder.")
            
    logger.info("Loader.py direct test finished.")